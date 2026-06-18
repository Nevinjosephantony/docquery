from docx import Document


def ingest_docx(filepath):
    print(f"Ingesting: {filepath}")
    return Document(filepath)


def get_avg_font_size(doc):
    sizes = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size:
                sizes.append(run.font.size / 12700)  # convert EMU to points
    return sum(sizes) / len(sizes) if sizes else 12  # default 12pt if nothing found


def is_fake_heading(para, avg_size):
    text = para.text.strip()
    if not text:
        return False

    is_short = len(text.split()) <= 8
    has_no_period = not text.strip().endswith(".")
    is_bold = any(run.bold for run in para.runs)

    size = para.runs[0].font.size if para.runs else None
    is_large = (size / 12700 > avg_size + 2) if size else False

    return (is_short and has_no_period) or is_bold or is_large


def chunk_by_fixed_size(doc, source_file, words_per_chunk=150):
    """Case 1 fallback — no headings at all, split by word count"""
    chunks = []
    chunk_index = 0
    current_text = ""

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        current_text += text + "\n"

        if len(current_text.split()) >= words_per_chunk:
            chunks.append({
                "chunk_id": f"chunk_{chunk_index}",
                "text": current_text.strip(),
                "metadata": {
                    "source": source_file,
                    "section": f"Section {chunk_index + 1}",
                    "parent_section": "Document",
                    "chunk_index": chunk_index
                }
            })
            chunk_index += 1
            current_text = ""

    # leftover text
    if current_text.strip():
        chunks.append({
            "chunk_id": f"chunk_{chunk_index}",
            "text": current_text.strip(),
            "metadata": {
                "source": source_file,
                "section": f"Section {chunk_index + 1}",
                "parent_section": "Document",
                "chunk_index": chunk_index
            }
        })

    return chunks


def chunk_docx(doc, source_file):
    chunks = []
    current_h1 = "Document"
    current_h2 = None
    current_text = ""
    chunk_index = 0
    avg_size = get_avg_font_size(doc)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name
        
        # Case 3 — proper heading styles
        if style == "Heading 1":
            if current_text.strip():
                chunks.append({
                    "chunk_id": f"chunk_{chunk_index}",
                    "text": current_text.strip(),
                    "metadata": {
                        "source": source_file,
                        "section": current_h2 or current_h1,
                        "parent_section": current_h1,
                        "chunk_index": chunk_index
                    }
                })
                chunk_index += 1
            current_h1 = text
            current_h2 = None
            current_text = ""

        elif style == "Heading 2":
            if current_text.strip():
                chunks.append({
                    "chunk_id": f"chunk_{chunk_index}",
                    "text": current_text.strip(),
                    "metadata": {
                        "source": source_file,
                        "section": current_h2 or current_h1,
                        "parent_section": current_h1,
                        "chunk_index": chunk_index
                    }
                })
                chunk_index += 1
            current_h2 = text
            current_text = ""

        # Case 2 — fake headings (bold, caps, large font)
        elif is_fake_heading(para, avg_size):
            if current_text.strip():
                chunks.append({
                    "chunk_id": f"chunk_{chunk_index}",
                    "text": current_text.strip(),
                    "metadata": {
                        "source": source_file,
                        "section": current_h2 or current_h1,
                        "parent_section": current_h1,
                        "chunk_index": chunk_index
                    }
                })
                chunk_index += 1
            # treat this paragraph as a heading
            current_h2 = text
            current_text = ""

        else:
            current_text += text + "\n"

    # save last chunk
    if current_text.strip():
        chunks.append({
            "chunk_id": f"chunk_{chunk_index}",
            "text": current_text.strip(),
            "metadata": {
                "source": source_file,
                "section": current_h2 or current_h1,
                "parent_section": current_h1,
                "chunk_index": chunk_index
            }
        })

    # Case 1 fallback — no headings found at all
    if len(chunks) <= 1:
        print("No headings detected, falling back to fixed-size chunking")
        return chunk_by_fixed_size(doc, source_file)

    return chunks